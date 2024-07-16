from PIL import Image 
import pytesseract 
import sys 
from pdf2image import convert_from_path 
from pytesseract import image_to_string
import os 

from ModuleArchivage.dao.dao_dossier import dao_dossier
from ModuleArchivage.dao.dao_document import dao_document
import os.path
from pprint import pprint

from django.conf import settings
import urllib

from django.core.files.storage import default_storage

import xlrd
import numpy as np

import docx2txt
import docx


class archive(object):

    @staticmethod
    def archiver(document, nom_document,type_document,reference_document,auteur,description = None,dossier_id = None, est_verifie = False, status = None):

        try:
            base_dir = settings.BASE_DIR
            media_dir = settings.MEDIA_ROOT
            media_url = settings.MEDIA_URL

            contenu = ""

            #print(nom_document)
            #print("Debut document ")
            
            docs_dir = 'documents/'
            media_dir = media_dir + '/' + docs_dir
            save_path = os.path.join(media_dir, str(nom_document))
            path = default_storage.save(save_path, document)
            url_document = media_url + docs_dir + str(nom_document)

            #print("PATH %s" % path)


            if nom_document[-3:] == "pdf":
                #print("PDF")
                pages = convert_from_path(path, 500) 
                image_counter = 1
                for page in pages: 
                    filename = "page_"+str(image_counter)+".jpg"
                    page.save(filename, 'JPEG') 
                    image_counter = image_counter + 1
                
                filelimit = image_counter-1

                for i in range(1, filelimit + 1):
                    filename = "page_"+str(i)+".jpg"
                    text = str(((pytesseract.image_to_string(Image.open(filename))))) 
                    text = text.replace('-\n', '')
                    #messages.error(request,text)
                    contenu += text
                
            elif nom_document[-3:] == "jpg" or nom_document[-3:] == "png" or nom_document[-3:] == "jpeg":
                #print("JPG")
                img=Image.open(path)
                contenu = image_to_string(img)

            elif nom_document[-3:] == "xls" or nom_document[-4:] == "xlsx":
                #print("EXCEL")
                workbook = xlrd.open_workbook(path)
                SheetNameList = workbook.sheet_names()
                for i in np.arange( len(SheetNameList) ):
                    a = 1

                worksheet = workbook.sheet_by_name(SheetNameList[0])
                num_rows = worksheet.nrows 
                num_cells = worksheet.ncols 
                
                curr_row = 0
                while curr_row < num_rows:
                    row = worksheet.row(curr_row)

                    curr_cell = 0
                    while curr_cell < num_cells:
                        
                        cell_type = worksheet.cell_type(curr_row, curr_cell)
                        cell_value = worksheet.cell_value(curr_row, curr_cell)
                        contenu += " " + str(cell_value)
                        curr_cell += 1
                    curr_row += 1

            elif nom_document[-3:] == "doc" or nom_document[-4:] == "docx":

                #print("WORD")
                doc = docx.Document(path)
 
                # read in each paragraph in file
                contenu = [p.text for p in doc.paragraphs]
                

            #print(contenu)
            document=dao_document.toCreateDocument(type_document,url_document,reference_document,description,est_verifie,status,contenu,dossier_id)
            document=dao_document.toSaveDocument(auteur, document)
            #document.index = contenu
            #document.save()

            return True

        except Exception as e:
            #print("Erreur d'archivage")
            #print(e)
            return False

